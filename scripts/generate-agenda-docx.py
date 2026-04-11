from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Style setup --
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0x7E)
    h.font.name = 'Calibri'

# -- Helper functions --
def add_colored_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7E)
    return h

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_body(text):
    p = doc.add_paragraph(text)
    return p

# ================================================================
# COVER / TITLE
# ================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AGENDA MEETING')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Presentasi Website Sea-Quill')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Tanggal: ').bold = True
info.add_run('9 April 2026\n')
info.add_run('Durasi: ').bold = True
info.add_run('± 90-120 menit\n')
info.add_run('URL: ').bold = True
info.add_run('https://seaquill-cms-production.up.railway.app/')

doc.add_page_break()

# ================================================================
# 1. OPENING
# ================================================================
add_colored_heading('1. Opening & Perkenalan (5 menit)', level=1)
add_bullet('Salam & perkenalan tim')
add_bullet('Tujuan meeting: review website yang sudah dibangun, diskusi konten, dan rencana go-live')

# ================================================================
# 2. OVERVIEW
# ================================================================
add_colored_heading('2. Overview Teknologi & Arsitektur (5 menit)', level=1)
add_bullet('Next.js 16 (React, server-side rendering)', bold_prefix='Framework: ')
add_bullet('PostgreSQL (semua konten dari database, bisa dikelola via admin)', bold_prefix='Database: ')
add_bullet('Railway (cloud hosting)', bold_prefix='Hosting: ')
add_bullet('JWT authentication, XSS protection, rate limiting, security headers', bold_prefix='Keamanan: ')
add_bullet('Sitemap otomatis, structured data Google, meta tags per halaman', bold_prefix='SEO: ')
add_bullet('Image optimization (AVIF/WebP), caching 60 detik, font optimization', bold_prefix='Performa: ')

# ================================================================
# 3. DEMO HALAMAN PUBLIK
# ================================================================
add_colored_heading('3. Demo Halaman Publik (40 menit)', level=1)

# 3.1 Global Layout
add_colored_heading('3.1 Global Layout (semua halaman)', level=2)
add_body('Setiap halaman memiliki:')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Preloader')
run.bold = True
p.add_run(' — Animasi loading saat halaman dibuka')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Header (sticky saat scroll):')
run.bold = True

items = [
    'Logo + badge info (bisa diganti dari admin)',
    'Navigasi utama: Beranda, Tentang, Produk (submenu kategori otomatis dari DB), Artikel (submenu kategori otomatis dari DB), Promo & Event, Galeri, Hubungi Kami',
    'Tombol marketplace (Tokopedia & Shopee)',
    'Menu hamburger untuk mobile',
    'Side drawer: logo, deskripsi, recent posts, social media',
    'Search popup (fullscreen)',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet 2')

p = doc.add_paragraph(style='List Bullet')
run = p.add_run('Footer (4 kolom):')
run.bold = True

items = [
    'Kolom 1: Logo + deskripsi perusahaan',
    'Kolom 2: Link produk per kategori (otomatis dari DB)',
    'Kolom 3: Link artikel per kategori (otomatis dari DB)',
    'Kolom 4: Bantuan & Support (FAQ, Kontak, Privasi, S&K)',
    'Info kontak: telepon, email, alamat',
    'Tagline + tombol marketplace (Tokopedia & Shopee)',
    'Copyright + ikon social media (Facebook, Twitter, LinkedIn, Instagram, WhatsApp)',
]
for item in items:
    p = doc.add_paragraph(item, style='List Bullet 2')

add_bullet(' button', bold_prefix='Scroll to Top')
add_bullet(' scroll (fade-in, slide effects)', bold_prefix='Animasi')

# 3.2 Homepage
add_colored_heading('3.2 Homepage /', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Hero Slider', 'Slider full-width dengan efek fade. Setiap slide: subtitle, judul utama, 2 tombol CTA ("Jelajahi Produk" & "Cek Manfaat Suplemen"), gambar background. Autoplay + pagination dots. Link social media di sisi kiri. Tombol scroll-down.'],
        ['2', 'Feature Cards', '3 kartu: "Tips & Edukasi Kesehatan", "Terstandar BPOM & Halal", "100% Original & Promo Menarik". Masing-masing punya judul, deskripsi, tombol CTA.'],
        ['3', 'About Section', 'Judul "Pilihan Tepat, untuk Hidup Sehat dan Berkualitas". Paragraf deskripsi perusahaan. Checklist 5 poin (BPOM, rekomendasi ahli, formula Indonesia, gaya hidup sehat, pilihan lengkap). Tombol "Selengkapnya Tentang Sea-Quill". 3 gambar dekoratif.'],
        ['4', 'Service List', 'Judul "Manfaat Utama Produk Sea-Quill". 4 layanan: Imunitas & Daya Tahan Tubuh, Kesehatan Jantung & Pembuluh Darah, Nutrisi Anak & Keluarga, Vitalitas & Energi. Masing-masing: ikon, judul, deskripsi, gambar, tombol link.'],
        ['5', 'Produk Unggulan', 'Slider produk (maks 12 produk dengan badge "Best Seller" atau "New"). Setiap kartu: gambar, badge, judul, deskripsi. Navigasi panah kiri/kanan. Autoplay.'],
        ['6', 'Brand Partners', 'Slider logo partner resmi: Kimia Farma, Watsons, Guardian, Wellings, Boston, Century, Boots. Grayscale → berwarna saat hover. Autoplay.'],
        ['7', 'Blog & Berita', 'Slider 6 artikel terbaru. Setiap kartu: gambar, nama penulis, judul, link "Selengkapnya". Navigasi panah.'],
    ]
)

# 3.3 Produk
add_colored_heading('3.3 Halaman Produk /produk', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner dengan judul halaman'],
        ['2', 'Kategori Produk', '8 kartu kategori (ikon, judul, deskripsi, tombol "Selengkapnya"): Beauty & Health, Mom & Kids, Multivitamin, Vitamin & Mineral, Kesehatan Jantung, Kesehatan Kulit, Kesehatan Sendi, Vitamin Kolesterol'],
        ['3', 'Grid Produk', 'Semua produk dalam grid. Filter pencarian client-side (ketik nama/tag). Setiap kartu: gambar + badge Best Seller/New, nama, tombol "LIHAT DETAIL", ikon aksi'],
        ['4', 'Pagination', 'Tombol halaman 1-4 + Next (UI only, belum fungsional)'],
    ]
)

# 3.4 Kategori Produk
add_colored_heading('3.4 Halaman Kategori Produk /produk/[kategori]', level=2)
add_bullet('Sama seperti halaman produk, tapi terfilter per kategori')
add_bullet('Mendukung kategori virtual: "Best Seller" dan "Produk Baru"')
add_bullet('Judul halaman otomatis sesuai nama kategori')

# 3.5 Detail Produk
add_colored_heading('3.5 Detail Produk /produk/detail/[slug]', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'Galeri Gambar', 'Gambar utama besar + strip thumbnail di bawah. Klik thumbnail untuk ganti gambar utama. Navigasi panah.'],
        ['3', 'Info Produk', 'Nama produk, rating bintang, jumlah review, deskripsi singkat, checklist fitur, tombol marketplace (Tokopedia & Shopee), SKU, kategori, tags, status stok'],
        ['4', 'Tab Konten', 'Tab 1: Deskripsi — Konten HTML lengkap. Tab 2: Informasi Tambahan — Komposisi, dosis, kategori, SKU. Tab 3: Ulasan — Review produk dengan rating bintang.'],
        ['5', 'Produk Terkait', 'Slider maks 6 produk dari kategori sama. Panah navigasi.'],
    ]
)

# 3.6 Artikel
add_colored_heading('3.6 Halaman Artikel /artikel', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'Grid Artikel (kolom 8/12)', 'Kartu artikel: gambar, nama penulis, judul, tombol "Read More"'],
        ['3', 'Sidebar (kolom 4/12)', 'Search — Kotak pencarian. Recent Posts — 4 artikel terbaru. Form Konsultasi Gratis — Input nama, email, telepon, pesan. Banner CTA — Tombol "Hubungi Kami"'],
        ['4', 'Pagination', 'UI only, belum fungsional'],
    ]
)

# 3.7 Kategori Artikel
add_colored_heading('3.7 Halaman Kategori Artikel /artikel/[kategori]', level=2)
add_bullet('Sama seperti halaman artikel, terfilter per kategori')
add_bullet('5 kategori: Tips Sehat, Imun & Daya Tahan Tubuh, Nutrisi & Gizi, Panduan Suplemen, Kecantikan')

# 3.8 Detail Artikel
add_colored_heading('3.8 Detail Artikel /artikel/detail/[slug]', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'Konten Artikel (8/12)', 'Gambar utama, meta (penulis, tanggal, kategori), judul, konten HTML lengkap, tags, link share social media (Facebook, Twitter, LinkedIn, Instagram)'],
        ['3', 'Komentar', '3 komentar sample (Rina, Budi, Siti) dengan balasan bersarang'],
        ['4', 'Form Komentar', 'Input nama, email, website, komentar + checkbox simpan data (UI only, belum terhubung)'],
        ['5', 'Sidebar (4/12)', 'Sama seperti halaman artikel listing'],
    ]
)

# 3.9 Tentang
add_colored_heading('3.9 Tentang Sea-Quill /tentang', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'About Area', 'Gambar, subtitle, heading, paragraf deskripsi, checklist poin-poin, gambar/video (mendukung mp4/webm), tombol play video'],
        ['3', 'Feature Cards', '3 kartu dengan judul, deskripsi, tombol CTA'],
        ['4', 'Project Slider', 'Slider dengan konten tentang perusahaan. Masing-masing: gambar, judul, teks. Autoplay + navigasi panah.'],
    ]
)

# 3.10 Galeri
add_colored_heading('3.10 Galeri /galeri', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'Slider per Platform', 'Satu slider per platform social media (Instagram, Facebook, TikTok). 1-4 item per view. Setiap item: gambar dengan overlay ikon platform, link ke posting di platform tersebut.'],
    ]
)

# 3.11 Hubungi Kami
add_colored_heading('3.11 Hubungi Kami /hubungi-kami', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'Social Media', 'Tab interface — klik platform (Instagram, Facebook, TikTok) untuk lihat deskripsi + tombol kunjungi. Animasi dekoratif.'],
        ['3', 'Marketplace', 'Tab interface — klik marketplace (Tokopedia, Shopee, dll) untuk lihat deskripsi + tombol beli.'],
        ['4', 'Form Kontak', 'Form lengkap: Nama, Email, Subject, Telepon, Pesan → kirim ke server. Mendukung custom form dari Form Builder. Loading state + pesan sukses/error. Opsi redirect ke WhatsApp.'],
    ]
)

# 3.12 Promo & Event
add_colored_heading('3.12 Promo & Event /promo', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', 'Banner judul halaman'],
        ['2', 'Promo', 'Grid kartu promo: gambar, judul, deskripsi, ikon panah'],
        ['3', 'Event', 'Event pertama tampil besar, selebihnya format kecil. Setiap event: tanggal, lokasi, judul, deskripsi, tombol CTA. Animasi dekoratif.'],
    ]
)

# 3.13 Pencarian
add_colored_heading('3.13 Pencarian /search', level=2)
add_table(
    ['#', 'Section', 'Isi'],
    [
        ['1', 'Breadcrumb', '"Hasil Pencarian"'],
        ['2', 'Judul', '"Hasil pencarian untuk \'{query}\' (X ditemukan)"'],
        ['3', 'Hasil Produk', 'Grid kartu produk yang cocok (maks 20 hasil)'],
        ['4', 'Hasil Artikel', 'Grid kartu artikel yang cocok (maks 20 hasil)'],
        ['5', 'Tidak Ada Hasil', 'Pesan "Tidak ada hasil" jika kosong'],
    ]
)
add_body('Pencarian mencari di: nama produk, deskripsi, tags + judul artikel, excerpt, tags. Minimal 2 karakter.')

# ================================================================
# 4. DEMO ADMIN PANEL
# ================================================================
add_colored_heading('4. Demo Admin Panel / CMS (20 menit)', level=1)
add_table(
    ['#', 'Menu', 'Fitur'],
    [
        ['1', 'Dashboard', 'Kartu statistik (produk, artikel, kategori, pesan). 3 chart (bar: produk per kategori, line: artikel per bulan, donut: pesan per bulan). Daftar pesan terbaru. Daftar produk terbaru.'],
        ['2', 'Produk', 'List + Create + Edit. Field: nama, slug, SKU, deskripsi singkat & lengkap (rich text), gambar, kategori, tags, harga, stok, badge (Best Seller/New), SEO fields, status draft/publish.'],
        ['3', 'Artikel', 'List + Create + Edit. Rich text editor untuk konten. Field: judul, slug, excerpt, author, kategori, tags, gambar, SEO fields, status draft/publish.'],
        ['4', 'Kategori Produk', 'CRUD 8 kategori dengan ikon, deskripsi, slug'],
        ['5', 'Kategori Artikel', 'CRUD 5 kategori blog'],
        ['6', 'Hero Slides', 'CRUD slider homepage (subtitle, judul, gambar background, CTA, urutan, aktif/nonaktif, group)'],
        ['7', 'Brand Partners', 'CRUD logo partner (nama, gambar, URL, urutan, aktif/nonaktif)'],
        ['8', 'Galeri', 'CRUD item galeri (gambar, caption, platform, URL, urutan)'],
        ['9', 'Promo', 'CRUD item promo (judul, deskripsi, gambar, aktif/nonaktif)'],
        ['10', 'Event', 'CRUD item event (judul, deskripsi, gambar, tanggal, lokasi, CTA)'],
        ['11', 'Pesan Masuk', 'Lihat & kelola pesan dari form kontak & konsultasi'],
        ['12', 'Users', 'Kelola akun admin (create/edit/delete)'],
        ['13', 'Media Library', 'Upload & browse semua file gambar/media'],
        ['14', 'Form Builder', 'Buat form custom dengan field bebas (text, email, select, radio, checkbox, dll). Preview form. Lihat submissions.'],
        ['15', 'Layout Pages', 'Editor per-halaman: Homepage, Tentang, Produk, Artikel, Galeri, Promo, Hubungi Kami — edit section tanpa coding'],
        ['16', 'Settings', 'Semua pengaturan situs: kontak, social media, deskripsi, copyright, tagline, footer navigasi, dll'],
        ['17', 'Activity Log', 'Histori semua perubahan (siapa edit/create/delete apa)'],
        ['18', 'Trash', 'Item yang sudah dihapus — bisa restore atau hapus permanen'],
    ]
)

# ================================================================
# 5. YANG PERLU DISIAPKAN
# ================================================================
add_colored_heading('5. Yang Perlu Disiapkan oleh Client (15 menit)', level=1)

# Prioritas Tinggi
add_colored_heading('Prioritas Tinggi (harus ada sebelum go-live)', level=2)
p = doc.add_paragraph()
run = p.add_run('🔴')
add_table(
    ['#', 'Item', 'Status Saat Ini', 'Yang Dibutuhkan'],
    [
        ['1', 'Artikel blog', '8 artikel placeholder Bahasa Inggris', 'Min 8-10 artikel Bahasa Indonesia tentang suplemen/kesehatan'],
        ['2', 'Nama penulis', 'Fiktif (Jane Doe, John Smith dll)', 'Nama Indonesia atau "Tim Redaksi Sea-Quill"'],
        ['3', 'Nomor telepon', 'Placeholder 021-1234-5678', 'Nomor asli'],
        ['4', 'WhatsApp', 'Placeholder 6281234567890', 'Nomor WhatsApp asli'],
        ['5', 'Email', 'info.seaquill@gmail.com', 'Konfirmasi atau ganti ke email resmi'],
        ['6', 'Alamat', 'Jl. Harmoni Plaza Blok A No. 8', 'Verifikasi alamat'],
    ]
)

# Prioritas Sedang
add_colored_heading('Prioritas Sedang', level=2)
p = doc.add_paragraph()
run = p.add_run('🟡')
add_table(
    ['#', 'Item', 'Status Saat Ini', 'Yang Dibutuhkan'],
    [
        ['7', 'Handle social media', 'Tidak konsisten antar halaman', 'Konfirmasi handle resmi (Instagram, Facebook, TikTok, YouTube, LinkedIn)'],
        ['8', 'Foto detail produk', '3 foto sama untuk semua produk', '3-4 foto unik per produk'],
        ['9', 'Halaman FAQ', 'Belum ada (404)', 'Konten FAQ'],
        ['10', 'Kebijakan Privasi', 'Belum ada (404)', 'Konten kebijakan privasi'],
        ['11', 'Syarat & Ketentuan', 'Belum ada (404)', 'Konten S&K'],
        ['12', 'Link Lazada & TikTok Shop', 'Kosong', 'URL toko resmi'],
    ]
)

# Opsional
add_colored_heading('Opsional', level=2)
p = doc.add_paragraph()
run = p.add_run('🟢')
add_table(
    ['#', 'Item', 'Keterangan'],
    [
        ['13', 'Foto blog', 'Saat ini stock/template, bisa diganti foto branded'],
        ['14', 'Video profil', 'Halaman Tentang mendukung video (mp4/webm)'],
        ['15', 'Event & promo aktif', 'Jika ada promo/event berjalan'],
    ]
)

# ================================================================
# 6. FEEDBACK
# ================================================================
add_colored_heading('6. Feedback & Diskusi (15 menit)', level=1)
add_bullet('Review desain & layout — ada yang ingin diubah?')
add_bullet('Warna, font, atau branding — sudah sesuai?')
add_bullet('Apakah ada halaman/fitur tambahan yang dibutuhkan?')
add_bullet('Prioritas revisi dari client')

# ================================================================
# 7. TIMELINE
# ================================================================
add_colored_heading('7. Timeline & Next Steps (10 menit)', level=1)
add_bullet('Deadline konten dari client')
add_bullet('Timeline revisi & testing')
add_bullet('Custom domain (pengganti subdomain Railway)')
add_bullet('Target tanggal go-live')
add_bullet('Rencana maintenance & update berkala')

# ================================================================
# 8. PENUTUP
# ================================================================
add_colored_heading('8. Penutup & Q&A (5 menit)', level=1)
add_bullet('Pertanyaan dari client')
add_bullet('Rangkuman action items')
add_bullet('Jadwal meeting berikutnya')

# ================================================================
# SAVE
# ================================================================
output_path = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'MEETING_AGENDA.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
